#!/usr/bin/env python3
"""
redline_to_docx.py — turn a review-state JSON `redline.edits` block into a Word
.docx with NATIVE tracked changes (w:ins / w:del) that opens in Word / Google
Docs as a real redline the attorney can accept or reject edit by edit.

TWO MODES
  1. TRUE inline redline (preferred): pass the ACTUAL letter as the source
     (`source.path` -> a .docx, or `source.text` -> the plain text). Each edit is
     located in the source by its `anchor` (falling back to `before`) and the
     change is recorded *in place*, so the reviewer reads the real document with
     the OLF markup on it.
  2. Clause-list fallback: when no source is available, the edits are rendered as
     a heading-per-provision list of tracked changes. Still a real redline, just
     detached from the surrounding text.

Nothing is sent. This writes a file to disk only — transmittal happens in the
lawyer's own client, never here.

Input JSON shape (review-state):
{
  "matter":  { "title": "...", "doc_family": "engagement_letter|commercial_contract",
               "sub_type": "...", "parties": {"client": "...", "counterparty": "..."} },
  "source":  { "path": "uploads/letter.docx" }  |  { "text": "full plain text..." }  |  null,
  "redline": { "edits": [
      { "op": "replace|insert|delete",
        "anchor": "short unique text that locates the spot (optional)",
        "before": "exact text to delete (replace/delete)",
        "after":  "text to insert (replace/insert)",
        "section": "clause label",
        "rationale": "one line why",
        "gate_class": "card|silent" }
  ] }
}

Usage:
    python tools/redline_to_docx.py review-state.json
    python tools/redline_to_docx.py review-state.json -o out.redline.docx

No third-party deps beyond python-docx (pip install python-docx).
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

# Deterministic revision metadata — the export is reproducible.
REV_DATE = "2026-08-14T00:00:00Z"
AUTHOR_CARD = "OLF Review (attorney decision)"
AUTHOR_SILENT = "OLF Review (administrative)"


# ---------------------------------------------------------------- primitives
def _plain_run(text):
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:rPr"))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _del_run(text):
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:rPr"))
    t = OxmlElement("w:delText")           # deleted text uses w:delText
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _ins_el(text, rev_id, author):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), REV_DATE)
    ins.append(_plain_run(text))
    return ins


def _del_el(text, rev_id, author):
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), str(rev_id))
    dele.set(qn("w:author"), author)
    dele.set(qn("w:date"), REV_DATE)
    dele.append(_del_run(text))
    return dele


def enable_track_changes(document):
    settings = document.settings.element
    if settings.find(qn("w:trackChanges")) is None:
        settings.append(OxmlElement("w:trackChanges"))


def _clear_paragraph_runs(paragraph):
    for r in list(paragraph._p.findall(qn("w:r"))):
        paragraph._p.remove(r)


# ------------------------------------------------------------- inline redline
class Rev:
    """Monotonic revision-id counter shared across the whole document."""
    def __init__(self):
        self.n = 0

    def next(self):
        self.n += 1
        return self.n


def _apply_inline(paragraph, edit, rev):
    """Rewrite `paragraph` so `before`/`after` become tracked changes in place.
    Returns True if the edit's anchor/before text was found and applied."""
    full = paragraph.text
    op = edit.get("op", "replace")
    before = edit.get("before") or ""
    after = edit.get("after") or ""
    anchor = edit.get("anchor") or before or after
    author = AUTHOR_CARD if edit.get("gate_class") == "card" else AUTHOR_SILENT

    # Locate the span to act on.
    needle = before if (op in ("replace", "delete") and before) else anchor
    if not needle or needle not in full:
        return False
    idx = full.index(needle)
    pre, post = full[:idx], full[idx + len(needle):]

    _clear_paragraph_runs(paragraph)
    p = paragraph._p
    if pre:
        p.append(_plain_run(pre))
    if op == "delete":
        p.append(_del_el(before, rev.next(), author))
    elif op == "insert":
        # keep the anchor, insert `after` immediately behind it
        p.append(_plain_run(needle))
        p.append(_ins_el(after, rev.next(), author))
    else:  # replace
        if before:
            p.append(_del_el(before, rev.next(), author))
        if after:
            p.append(_ins_el(after, rev.next(), author))
    if post:
        p.append(_plain_run(post))
    return True


def _source_document(src):
    """Return (Document, made_from_text) for the source letter, or (None, False)."""
    if not src:
        return None, False
    path = src.get("path")
    if path and os.path.exists(path) and path.lower().endswith(".docx"):
        return Document(path), False
    text = src.get("text")
    if text:
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(11)
        for line in text.splitlines():
            doc.add_paragraph(line)
        return doc, True
    return None, False


def build_inline(m, out_path):
    src = m.get("source")
    doc, _ = _source_document(src)
    if doc is None:
        return None  # signal caller to use fallback
    enable_track_changes(doc)
    rev = Rev()
    edits = ((m.get("redline") or {}).get("edits")) or []
    applied, missed = 0, []
    for e in edits:
        hit = False
        for para in doc.paragraphs:
            if _apply_inline(para, e, rev):
                hit = True
                applied += 1
                break
        if not hit:
            missed.append(e)

    # Any edit whose anchor was not found is appended as an addendum so nothing
    # is silently dropped.
    if missed:
        doc.add_page_break()
        h = doc.add_heading("OLF markup — edits not anchored inline", level=1)
        note = doc.add_paragraph()
        nr = note.add_run(
            "The anchor text for the changes below was not located verbatim in "
            "the source. Placement is a judgment call for the reviewer."
        )
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _render_clause_list(doc, missed, rev)

    doc.save(out_path)
    return {"applied": applied, "missed": len(missed), "mode": "inline"}


# -------------------------------------------------------------- clause list
def _render_clause_list(doc, edits, rev):
    for e in edits:
        author = AUTHOR_CARD if e.get("gate_class") == "card" else AUTHOR_SILENT
        doc.add_heading(e.get("section", "Provision"), level=2)
        gc = e.get("gate_class", "card")
        meta = doc.add_paragraph()
        mr = meta.add_run(
            "Attorney decision required" if gc == "card" else "Administrative / auto-applied"
        )
        mr.font.size = Pt(8)
        mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        p = doc.add_paragraph()
        op = e.get("op", "replace")
        before = e.get("before") or ""
        after = e.get("after") or ""
        if op in ("replace", "delete") and before:
            p._p.append(_del_el(before, rev.next(), author))
        if op in ("replace", "insert") and after:
            p._p.append(_ins_el(after, rev.next(), author))

        if e.get("rationale"):
            why = doc.add_paragraph()
            wr = why.add_run(f"Why: {e['rationale']}")
            wr.italic = True
            wr.font.size = Pt(9)
            wr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph()


def build_fallback(m, out_path):
    doc = Document()
    enable_track_changes(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    matter = m.get("matter") or {}
    fam = (matter.get("doc_family") or "Document").replace("_", " ").title()
    sub = (matter.get("sub_type") or "").replace("_", " ")
    parties = matter.get("parties") or {}
    who = parties.get("counterparty") or parties.get("client") or ""
    title = fam + (f" ({sub})" if sub else "") + (f" — {who}" if who else "")
    doc.add_heading(title, level=0)
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "OLF proposed changes, recorded as tracked changes. No source document "
        "was supplied, so edits are listed by provision rather than marked in "
        "place. Accept or reject in Word; nothing has been sent."
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    edits = ((m.get("redline") or {}).get("edits")) or []
    if not edits:
        doc.add_paragraph("No redline was produced for this matter.")
        doc.save(out_path)
        return {"applied": 0, "missed": 0, "mode": "empty"}
    _render_clause_list(doc, edits, Rev())
    doc.save(out_path)
    return {"applied": len(edits), "missed": 0, "mode": "clause-list"}


def build(m, out_path):
    result = build_inline(m, out_path)
    if result is None:
        result = build_fallback(m, out_path)
    return result


def default_out(in_path):
    base = os.path.basename(in_path).replace(".review-state.json", "").replace(".json", "")
    d = os.path.dirname(in_path) or "."
    return os.path.join(d, f"{base}.redline.docx")


def main():
    ap = argparse.ArgumentParser(description="Export a review-state redline to a tracked-changes .docx")
    ap.add_argument("input", help="review-state JSON file")
    ap.add_argument("-o", "--output", help="output .docx path")
    args = ap.parse_args()
    with open(args.input) as f:
        m = json.load(f)
    out_path = args.output or default_out(args.input)
    res = build(m, out_path)
    print(f"{args.input}  ->  {out_path}   ({res['mode']}: {res['applied']} applied, {res['missed']} unanchored)")


if __name__ == "__main__":
    main()
